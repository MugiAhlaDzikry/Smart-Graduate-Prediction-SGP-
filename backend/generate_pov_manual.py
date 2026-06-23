from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_pov_manual():
    doc = Document()
    
    # ─── Judul Dokumen ───
    title = doc.add_heading('User Manual SGP - Edisi Presentasi', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Panduan Menjelaskan Tiap Halaman Aplikasi (POV Presenter)')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_page_break()

    # ─── Intro ───
    doc.add_heading('1. Pembukaan (Halaman Login)', level=1)
    doc.add_paragraph(
        '"Selamat pagi/siang Bapak, Ibu, dan teman-teman sekalian. Pada kesempatan kali ini, saya akan mendemonstrasikan '
        'langsung aplikasi Smart Graduate Predictor (SGP). Seperti yang kita lihat di layar, ini adalah halaman Login. '
        'Sistem ini telah dilengkapi dengan arsitektur Multi-Tenancy. Artinya, saat saya login dengan akun saya, semua data, '
        'model AI, dan riwayat prediksi yang muncul nantinya akan 100% terisolasi dan khusus untuk akun saya saja. Mari kita masuk ke aplikasinya..."'
    )

    # ─── Data Mahasiswa ───
    doc.add_heading('2. Halaman Data Mahasiswa', level=1)
    doc.add_paragraph(
        '"Hal pertama yang harus kita lakukan sebagai pengguna baru adalah menyiapkan data. Mari kita beralih ke halaman \'Data Mahasiswa\'."'
    )
    doc.add_paragraph('Poin yang dijelaskan saat mendemonstrasikan halaman ini:', style='List Bullet')
    doc.add_paragraph('"Di halaman ini, kita dapat mengunggah (upload) data historis mahasiswa dari tahun-tahun sebelumnya dalam bentuk Excel atau CSV."', style='List Bullet')
    doc.add_paragraph('"Sistem akan langsung menampilkan data tersebut dalam bentuk tabel yang mudah dibaca. Kita juga memiliki kendali penuh di sini; jika terjadi kesalahan atau kita ingin mengulang dari awal, kita cukup menekan tombol merah \'Hapus Semua Data\' yang otomatis akan mereset dataset dan model kita."', style='List Bullet')

    # ─── Training Model ───
    doc.add_heading('3. Halaman Training Model', level=1)
    doc.add_paragraph(
        '"Setelah kita memiliki data, saatnya kita \'mengajari\' kecerdasan buatan (AI) kita. Mari kita buka halaman \'Training Model\'."'
    )
    doc.add_paragraph('Poin yang dijelaskan saat mendemonstrasikan halaman ini:', style='List Bullet')
    doc.add_paragraph('"Di sinilah keajaiban Machine Learning terjadi. Kita cukup menekan tombol \'Mulai Training\', dan algoritma XGBoost akan menganalisis pola dari ratusan data mahasiswa yang tadi kita unggah."', style='List Bullet')
    doc.add_paragraph('"Proses ini sangat transparan. Setelah selesai, kita bisa langsung melihat seberapa akurat AI kita. Sebagai contoh di layar, akurasi model kita mencapai angka di atas 90%, lengkap dengan metrik Presisi dan Recall-nya."', style='List Bullet')

    # ─── Dashboard ───
    doc.add_heading('4. Halaman Dashboard', level=1)
    doc.add_paragraph(
        '"Sekarang, mari kita lihat halaman utama kita, yaitu \'Dashboard\'."'
    )
    doc.add_paragraph('Poin yang dijelaskan saat mendemonstrasikan halaman ini:', style='List Bullet')
    doc.add_paragraph('"Dashboard ini berfungsi sebagai pusat kendali ringkasan. Di sini kita langsung disuguhkan informasi penting, seperti total populasi mahasiswa kita, berapa persen yang diprediksi lulus tepat waktu, dan yang paling krusial: berapa banyak mahasiswa yang masuk kategori risiko tinggi."', style='List Bullet')
    doc.add_paragraph('"Ini memberikan insight instan bagi Program Studi untuk segera mengambil tindakan intervensi bagi mahasiswa yang berisiko."', style='List Bullet')

    # ─── Prediksi Kelulusan ───
    doc.add_heading('5. Halaman Prediksi Kelulusan', level=1)
    doc.add_paragraph(
        '"Inilah inti dari aplikasi ini. Mari kita coba menggunakan AI yang sudah kita latih tadi di halaman \'Prediksi Kelulusan\'."'
    )
    doc.add_paragraph('Poin yang dijelaskan saat mendemonstrasikan halaman ini:', style='List Bullet')
    doc.add_paragraph('"Sistem menyediakan dua cara prediksi. Pertama, Prediksi Batch untuk memprediksi data puluhan mahasiswa baru sekaligus cukup dengan mengupload file. Sangat efisien!"', style='List Bullet')
    doc.add_paragraph('"Kedua, Prediksi Individu. Mari kita simulasikan langsung. Saya akan memasukkan seorang mahasiswa dengan IPK rendah dan kehadiran di bawah standar. Mari kita lihat reaksinya..."', style='List Bullet')
    doc.add_paragraph('"Dan lihat! AI bukan hanya menebak \'Terlambat Lulus\', tetapi aplikasi ini juga menampilkan \'Faktor Risiko\'. Kita diberitahu bahwa faktor utama penyebabnya adalah IPK dan Tingkat Kehadiran. Ini yang disebut Explainable AI, jadi keputusan AI ini sangat masuk akal dan transparan."', style='List Bullet')

    # ─── Riwayat Prediksi ───
    doc.add_heading('6. Halaman Riwayat Prediksi', level=1)
    doc.add_paragraph(
        '"Terakhir, kita bisa memantau semua rekam jejak tebakan sistem di halaman \'Riwayat Prediksi\'."'
    )
    doc.add_paragraph('Poin yang dijelaskan saat mendemonstrasikan halaman ini:', style='List Bullet')
    doc.add_paragraph('"Aplikasi secara otomatis mencatat setiap prediksi yang pernah kita lakukan beserta persentase keyakinannya (probabilitas). Ini berfungsi sebagai audit log sekaligus mempermudah Dosen Wali untuk meninjau kembali mahasiswa mana saja yang sebelumnya sudah dievaluasi."', style='List Bullet')

    # Output file
    output_path = os.path.join(os.path.dirname(__file__), 'User_Manual_POV_Presentasi.docx')
    doc.save(output_path)
    print(f"Berhasil membuat dokumen presentasi di: {output_path}")

if __name__ == '__main__':
    create_pov_manual()
