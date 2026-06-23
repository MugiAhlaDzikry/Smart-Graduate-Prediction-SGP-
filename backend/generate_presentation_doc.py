from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_presentation_doc():
    doc = Document()
    
    # ─── Judul Dokumen ───
    title = doc.add_heading('Materi Presentasi: Smart Graduate Predictor (SGP)', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Draft Konten Slide Presentasi Kelas')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_page_break()

    # ─── Slide 1 ───
    doc.add_heading('Slide 1: Judul Presentasi', level=1)
    p = doc.add_paragraph()
    p.add_run('SMART GRADUATE PREDICTOR (SGP)\n').bold = True
    p.add_run('Sistem Prediksi Kelulusan Mahasiswa Berbasis Machine Learning dengan Arsitektur Multi-Tenancy.\n')
    doc.add_paragraph('Poin Bicara:', style='List Bullet')
    doc.add_paragraph('Perkenalkan nama dan tim/kelompok.', style='List Bullet')
    doc.add_paragraph('Jelaskan secara singkat bahwa aplikasi ini dirancang untuk memprediksi apakah seorang mahasiswa akan lulus tepat waktu atau tidak berdasarkan data historis.', style='List Bullet')

    # ─── Slide 2 ───
    doc.add_heading('Slide 2: Latar Belakang & Masalah', level=1)
    doc.add_paragraph('Konten Slide:', style='List Bullet')
    doc.add_paragraph('Pihak kampus sering kesulitan mendeteksi dini mahasiswa yang berisiko lulus terlambat atau drop out.', style='List Bullet')
    doc.add_paragraph('Banyaknya data akademik yang menganggur dan tidak dimanfaatkan untuk pengambilan keputusan strategis.', style='List Bullet')
    doc.add_paragraph('Poin Bicara:', style='List Bullet')
    doc.add_paragraph('Sebutkan bahwa dengan adanya SGP, prodi atau dosen wali bisa melakukan intervensi (bimbingan ekstra) lebih awal kepada mahasiswa yang berisiko.', style='List Bullet')

    # ─── Slide 3 ───
    doc.add_heading('Slide 3: Solusi yang Ditawarkan', level=1)
    doc.add_paragraph('Konten Slide:', style='List Bullet')
    doc.add_paragraph('Memanfaatkan Algoritma Machine Learning (XGBoost/Random Forest).', style='List Bullet')
    doc.add_paragraph('Fitur utama: Dashboard Statistik, Prediksi Individu, Prediksi Batch, dan Training Custom Model.', style='List Bullet')
    doc.add_paragraph('Arsitektur Multi-Tenancy: Keamanan dan isolasi data untuk setiap pengguna.', style='List Bullet')
    doc.add_paragraph('Poin Bicara:', style='List Bullet')
    doc.add_paragraph('Jelaskan bahwa SGP bukan hanya sekadar menampilkan data, tapi menggunakan "AI" untuk belajar dari data lama (Training) dan menebak masa depan (Predicting).', style='List Bullet')

    # ─── Slide 4 ───
    doc.add_heading('Slide 4: Cara Kerja Sistem (Workflow)', level=1)
    doc.add_paragraph('1. Upload Data Mahasiswa Lama (Data Historis).')
    doc.add_paragraph('2. Latih Model ML (Sistem akan membuat Model AI khusus untuk user tersebut).')
    doc.add_paragraph('3. Prediksi Mahasiswa Baru (Individu / Upload file CSV secara massal).')
    doc.add_paragraph('4. Analisis Faktor Risiko (Menampilkan alasan mengapa AI memprediksi mahasiswa tersebut berisiko terlambat lulus).')

    # ─── Slide 5 ───
    doc.add_heading('Slide 5: Keunggulan Teknis', level=1)
    doc.add_paragraph('Konten Slide:', style='List Bullet')
    doc.add_paragraph('Multi-Tenancy: Setiap Dosen/Admin memiliki ruang kerjanya sendiri. Model AI milik Dosen A dilatih dari datanya sendiri dan tidak bercampur dengan Dosen B.', style='List Bullet')
    doc.add_paragraph('Explainable AI (SHAP): Tidak seperti black-box, AI kami bisa menjelaskan "Kenapa" (misal: "Diprediksi terlambat karena IPK < 2.75 dan Kehadiran < 70%").', style='List Bullet')
    doc.add_paragraph('Modern Tech Stack: ReactJS (Frontend), FastAPI (Backend), Supabase (Database/Auth).', style='List Bullet')

    # ─── Slide 6 ───
    doc.add_heading('Slide 6: Demo Aplikasi (Live Demo)', level=1)
    doc.add_paragraph('Poin Bicara & Urutan Demo:', style='List Bullet')
    doc.add_paragraph('1. Tunjukkan halaman Login/Daftar.', style='List Bullet')
    doc.add_paragraph('2. Tunjukkan menu Data Mahasiswa (Upload File Batch yang sudah disiapkan).', style='List Bullet')
    doc.add_paragraph('3. Lakukan Training Model di depan kelas, pamerkan Akurasi (misalnya 93%).', style='List Bullet')
    doc.add_paragraph('4. Masuk ke Prediksi Batch, upload mahasiswa baru, dan tunjukkan tabel hasil prediksinya (Lulus Tepat Waktu / Terlambat).', style='List Bullet')
    doc.add_paragraph('5. Buka Prediksi Individu, masukkan angka palsu (IPK 1.5) untuk menunjukkan bahwa AI berani memprediksi "Terlambat".', style='List Bullet')

    # ─── Slide 7 ───
    doc.add_heading('Slide 7: Kesimpulan & Penutup', level=1)
    doc.add_paragraph('Konten Slide:', style='List Bullet')
    doc.add_paragraph('Smart Graduate Predictor adalah solusi cerdas untuk mengawal masa studi mahasiswa secara proaktif.', style='List Bullet')
    doc.add_paragraph('Terima Kasih / Q&A.', style='List Bullet')
    doc.add_paragraph('Poin Bicara:', style='List Bullet')
    doc.add_paragraph('Sampaikan harapan bahwa aplikasi ini bisa dikembangkan lebih lanjut menjadi platform B2B untuk berbagai Universitas.', style='List Bullet')

    # Output file
    output_path = os.path.join(os.path.dirname(__file__), 'Materi_Presentasi_SGP.docx')
    doc.save(output_path)
    print(f"Berhasil membuat dokumen presentasi di: {output_path}")

if __name__ == '__main__':
    create_presentation_doc()
