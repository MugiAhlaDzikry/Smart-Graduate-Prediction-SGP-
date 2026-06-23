from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('User Manual', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Aplikasi Smart Graduate Predictor (SGP)')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(16)
    run.bold = True
    
    doc.add_page_break()

    # Section 1
    doc.add_heading('1. Pendahuluan', level=1)
    doc.add_paragraph(
        "Aplikasi Smart Graduate Predictor (SGP) adalah platform berbasis kecerdasan buatan "
        "(AI) yang dirancang untuk memprediksi apakah seorang mahasiswa akan lulus tepat waktu "
        "berdasarkan data akademik dan non-akademik. Sistem ini mendukung arsitektur multi-tenancy "
        "sehingga setiap pengguna (dosen/admin) memiliki dataset dan model prediksi yang sepenuhnya terisolasi."
    )

    # Section 2
    doc.add_heading('2. Navigasi dan Menu Utama', level=1)
    
    doc.add_heading('2.1. Dashboard', level=2)
    doc.add_paragraph(
        "Halaman utama yang memberikan ringkasan statistik (Overview). Anda dapat melihat total mahasiswa, "
        "persentase mahasiswa yang diprediksi lulus tepat waktu, mahasiswa dengan risiko tinggi, serta "
        "akurasi model AI saat ini. Dashboard akan kosong apabila Anda belum mengunggah data."
    )
    
    doc.add_heading('2.2. Prediksi Kelulusan', level=2)
    doc.add_paragraph(
        "Menu ini digunakan untuk melakukan prediksi terhadap data mahasiswa. Terdapat dua metode:"
    )
    doc.add_paragraph("• Prediksi Individu: Memasukkan data satu mahasiswa secara manual ke dalam form.", style='List Bullet')
    doc.add_paragraph("• Prediksi Batch: Mengunggah file (CSV/Excel) berisi data banyak mahasiswa sekaligus untuk diprediksi massal.", style='List Bullet')

    doc.add_heading('2.3. Riwayat Prediksi', level=2)
    doc.add_paragraph(
        "Menyimpan semua catatan hasil prediksi masa lalu yang pernah Anda lakukan, lengkap dengan probabilitas kelulusan "
        "dan faktor risikonya. Anda dapat menghapus riwayat prediksi di halaman ini."
    )

    doc.add_heading('2.4. Data Mahasiswa', level=2)
    doc.add_paragraph(
        "Halaman untuk mengelola dataset asli. Anda wajib mengunggah data historis mahasiswa Anda ke sini sebelum "
        "dapat melatih model AI. Terdapat tombol 'Hapus Semua Data' jika Anda ingin mengosongkan dataset dan melakukan reset."
    )

    doc.add_heading('2.5. Training Model', level=2)
    doc.add_paragraph(
        "Halaman untuk melatih ulang (retrain) model Machine Learning menggunakan dataset terbaru yang Anda unggah "
        "di halaman Data Mahasiswa. Menampilkan metrik akurasi, presisi, recall, dan F1 Score dari hasil training."
    )

    # Section 3
    doc.add_heading('3. Alur Penggunaan (Workflow)', level=1)
    
    workflow = [
        ("Langkah 1: Unggah Dataset", "Masuk ke menu 'Data Mahasiswa' dan unggah file CSV atau Excel yang berisi data historis mahasiswa (termasuk kolom 'is_on_time' yang valid)."),
        ("Langkah 2: Latih Model (Training)", "Buka menu 'Training Model' dan klik 'Mulai Training'. Tunggu hingga proses selesai dan metrik akurasi model Anda muncul."),
        ("Langkah 3: Lakukan Prediksi", "Masuk ke menu 'Prediksi Kelulusan' lalu pilih Individu atau Batch untuk mulai memprediksi mahasiswa baru menggunakan model yang telah Anda latih."),
        ("Langkah 4: Analisis Hasil", "Lihat hasil prediksi, termasuk persentase keyakinan model dan faktor-faktor spesifik yang paling mempengaruhi prediksi mahasiswa tersebut."),
        ("Langkah 5: Reset Sistem (Opsional)", "Jika Anda ingin mengganti dataset atau mereset model, Anda dapat kembali ke menu 'Data Mahasiswa' lalu mengklik 'Hapus Semua Data'. Hal ini akan turut mereset model ML Anda.")
    ]
    
    for title, desc in workflow:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    # Output file
    output_path = os.path.join(os.path.dirname(__file__), 'User_Manual_SGP.docx')
    doc.save(output_path)
    print(f"Berhasil membuat dokumen Word di: {output_path}")

if __name__ == '__main__':
    create_manual()
